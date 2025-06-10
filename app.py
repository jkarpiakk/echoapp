from flask import Flask, render_template, request, send_file
from io import BytesIO
import csv
from docx import Document
from datetime import datetime

app = Flask(__name__)

# Form fields and labels
field_defs = [
    ('patientName', 'Imię i nazwisko'),
    ('patientID', 'Nr badania / ID'),
    ('examDate', 'Data badania'),
    ('operator', 'Operator / lekarz'),
    ('LVIDd', 'LV (LVIDd/s)'),
    ('IVSd', 'IVSd'),
    ('PWd', 'PWd'),
    ('LA', 'LA'),
    ('LAVI', 'LAVI'),
    ('RAA', 'RAA'),
    ('RV', 'RV'),
    ('OpuszkaAorty', 'Opuszka aorty'),
    ('AortaWstepujaca', 'Aorta wstępująca'),
    ('MPA', 'Pień płucny (MPA)'),
    ('EF', 'Frakcja wyrzutowa (EF)'),
    ('EFmethod', 'Metoda EF'),
    ('EDV', 'EDV'),
    ('ESV', 'ESV'),
    ('segmentalPresent', 'Odcinkowe zaburzenia obecne'),
    ('segmentalDescription', 'Opis odcinkowych zaburzeń'),
    ('TAPSE', 'TAPSE'),
    ('mitralRegurg', 'Mitral regurgitacja'),
    ('mitralVC', 'Mitral VC'),
    ('mitralPISA', 'Mitral PISA'),
    ('mitralERO', 'Mitral ERO'),
    ('mitralEAratio', 'Mitral E/A ratio'),
    ('mitralMxPG', 'Mitral MxPG'),
    ('mitralMnPG', 'Mitral MnPG'),
    ('aorticMorphology', 'Aortal morfologia'),
    ('aorticRegSten', 'Aortal stenoza/regurgitacja'),
    ('aorticVC', 'Aortal VC'),
    ('aorticPHT', 'Aortal PHT'),
    ('aorticVmax', 'Aortal Vmax'),
    ('aorticPmax', 'Aortal Pmax'),
    ('tricuspidRegurg', 'Tricuspid regurgitacja'),
    ('tricuspidVmax', 'Tricuspid Vmax'),
    ('tricuspidPmax', 'Tricuspid Pmax'),
    ('pericardEffusion', 'Płyn w osierdziu'),
    ('pericardNotes', 'Uwagi osierdziowe'),
    ('remarks', 'Uwagi końcowe')
]

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/export', methods=['POST'])
def export():
    form = request.form.to_dict()
    fmt = form.pop('format', 'csv')
    data = {k: v for k, v in form.items() if v}

    if fmt == 'csv':
        buffer = BytesIO()
        writer = csv.writer(buffer)
        writer.writerow([label for _, label in field_defs])
        writer.writerow([data.get(key, '') for key, _ in field_defs])
        buffer.seek(0)
        filename = f"echo_{data.get('patientID','record')}.csv"
        return send_file(buffer, as_attachment=True, attachment_filename=filename, mimetype='text/csv')
    else:
        doc = Document()
        doc.add_heading('Formularz badania echokardiograficznego', level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parametr'
        hdr_cells[1].text = 'Wartość'
        for key, label in field_defs:
            if key in data:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = data[key]
        doc.add_paragraph(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"echo_{data.get('patientID','form')}.docx"
        return send_file(buffer, as_attachment=True, attachment_filename=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)
